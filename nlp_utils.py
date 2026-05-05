import PyPDF2
try:
    import docx
except ImportError:
    pass
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re

def extract_text_from_pdf(pdf_path):
    """Extracts text from a given PDF file."""
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + " "
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extracts text from a given DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_path)
        # Extract from normal paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + " "
        # Extract from tables (many resumes use tables for structural layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            text += para.text + " "
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text.strip()

def extract_text(file_path):
    """Router for file extraction based on extension."""
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    return ""

def preprocess_text(text):
    """Basic text cleaning."""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'\W+', ' ', text)
    return text.strip()

def extract_skills(text):
    """
    Extracts high-quality keywords (skills, technologies, certifications) mimicking an ATS scanner.
    """
    try:
        import spacy
        import re
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text)
        
        keywords = set()
        
        # 1. Predefined Master Skill Dictionary (Guaranteed hits for top tech)
        master_skills = [
            "Python", "Java", "JavaScript", "C++", "C#", "Ruby", "PHP", "Swift", "Kotlin", "Go", "Rust", "TypeScript",
            "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Cassandra", "Oracle", "DynamoDB",
            "AWS", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins",
            "React", "Angular", "Vue", "Node.js", "Django", "Flask", "Spring Boot", "Express", "Ruby on Rails",
            "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "TensorFlow", "PyTorch", "Scikit-Learn",
            "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Excel",
            "Agile", "Scrum", "Git", "GitHub", "GitLab", "CI/CD", "REST API", "GraphQL", "Microservices",
            "Linux", "Unix", "Bash", "Shell", "HTML", "CSS", "Tailwind"
        ]
        
        lower_text = text.lower()
        for skill in master_skills:
            escaped_skill = re.escape(skill.lower())
            # Use negative lookbehind/lookahead for word boundaries including symbols like C++
            if re.search(r'(?<![a-zA-Z])' + escaped_skill + r'(?![a-zA-Z])', lower_text):
                keywords.add(skill)

        # 2. Grab Entities (Organizations, Products often correlate to technologies/skills)
        for ent in doc.ents:
            if ent.label_ in ['ORG', 'PRODUCT']:
                # Filter out pure digits
                clean_ent = re.sub(r'[^a-zA-Z0-9\.\+\- ]', '', ent.text).strip().title()
                if len(clean_ent) > 2 and not clean_ent.isdigit():
                    keywords.add(clean_ent)
                    
        # 3. Grab technical-looking tokens (camelCase, contains dots/pluses like OpenCV, Node.js)
        for token in doc:
            clean_tok = token.text.strip()
            if ('+' in clean_tok or '.' in clean_tok or clean_tok.isupper()) and len(clean_tok) >= 2:
                if not token.is_stop and not clean_tok.isdigit():
                    keywords.add(clean_tok.title())
                    
        # 4. Grab Proper Nouns that are capitalized
        for token in doc:
            if token.pos_ == 'PROPN' and token.text.istitle() and not token.is_stop:
                if len(token.text) > 2:
                    keywords.add(token.text.strip().title())
                    
        # Filter out noise
        junk_words = {"University", "College", "School", "Inc", "Llc", "Company", "January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December", "State", "City", "Music", "Families", "Saadhana"}
        final_skills = [k for k in keywords if not any(j.lower() in k.lower() for j in junk_words)]
        
        return list(set(final_skills))[:30]
    except Exception as e:
        print(f"Extraction error: {e}")
        return []

import difflib

def calculate_similarity(job_description, resumes_text):
    """
    Calculates the cosine similarity between the Job Description and a list of resumes.
    Returns: list of similarity percentages
    """
    if not job_description or not resumes_text:
        return [0] * len(resumes_text)

    documents = [preprocess_text(job_description)] + [preprocess_text(txt) for txt in resumes_text]
    
    vectorizer = TfidfVectorizer(stop_words='english')
    try:
        tfidf_matrix = vectorizer.fit_transform(documents)
        cosine_similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
        return [round(score * 100, 2) for score in cosine_similarities]
    except ValueError:
        return [0] * len(resumes_text)

def generate_ats_metrics(resume_text, job_desc_text, semantic_score):
    """
    Generates advanced ATS metrics including actionable feedback and JD keyword matching.
    """
    metrics = {
        'semantic_score': semantic_score,
        'email': None,
        'phone': None,
        'word_count': 0,
        'action_verbs_found': 0,
        'quantifiable_metrics': 0,
        'missing_skills': [],
        'matched_skills': [],
        'feedback_tips': [],
        'ats_score': 0
    }
    
    # 1. Contact Info Extraction
    email_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', resume_text)
    if email_match:
        metrics['email'] = email_match.group(0)
        
    phone_match = re.search(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', resume_text)
    if phone_match:
        metrics['phone'] = phone_match.group(0)
        
    # 2. Word Count & Quantifiable Metrics
    words = resume_text.split()
    metrics['word_count'] = len(words)
    
    # Count numbers and percentages (recruiter impact)
    quant_metrics = len(re.findall(r'\b\d+%\b|\$\d+|\b\d+\s*(?:million|thousand|k|m)\b', resume_text.lower()))
    metrics['quantifiable_metrics'] = quant_metrics
    
    # 3. Action Verbs Count
    action_verbs = ['managed', 'developed', 'architected', 'led', 'designed', 'created', 'implemented', 'optimized', 'spearheaded', 'resolved', 'analyzed', 'driven', 'orchestrated', 'reduced', 'increased', 'maximized']
    lower_text = resume_text.lower()
    found_verbs = sum(1 for verb in action_verbs if verb in lower_text)
    metrics['action_verbs_found'] = min(found_verbs, 10)
    
    # 4. Strict JD Keyword Matching Jobscan-style using DiffLib Fuzzy
    jd_skills = extract_skills(job_desc_text)
    resume_skills = extract_skills(resume_text)
    
    matched = []
    missing = []
    
    valid_resume_skills = set([s.lower() for s in resume_skills])
    
    for jd_skill in jd_skills:
        jd_lower = jd_skill.lower()
        # Find best match in resume skills with 80% similarity threshold
        matches = difflib.get_close_matches(jd_lower, valid_resume_skills, n=1, cutoff=0.8)
        
        # Fallback to direct raw text search if extracted slightly wrong
        if matches or re.search(r'(?<![a-zA-Z])' + re.escape(jd_lower) + r'(?![a-zA-Z])', lower_text):
            matched.append(jd_skill)
        else:
            missing.append(jd_skill)
            
    metrics['missing_skills'] = missing[:12]
    metrics['matched_skills'] = matched[:12]
    
    keyword_match_pct = 100 if not jd_skills else (len(matched) / len(jd_skills)) * 100
    
    # Structural Section Analysis
    lower_text_clean = re.sub(r'[^a-zA-Z0-9]', '', lower_text)
    metrics['sections'] = {
        'education': 'education' in lower_text_clean or 'academic' in lower_text_clean,
        'experience': 'experience' in lower_text_clean or 'employment' in lower_text_clean,
        'projects': 'project' in lower_text_clean or 'portfolio' in lower_text_clean
    }
    
    # 5. Composite ATS Score (Jobscan style weighting)
    # 30% Semantic, 30% Keywords, 20% Impact (Quant/Verbs), 20% Formatting (Sections/Length)
    score = (semantic_score * 0.3) + (keyword_match_pct * 0.3)
    
    impact_score = 0
    if metrics['quantifiable_metrics'] > 3:
        impact_score += 50
    elif metrics['quantifiable_metrics'] > 0:
        impact_score += 25
    impact_score += (metrics['action_verbs_found'] / 10) * 50
    score += (impact_score * 0.2)
    
    format_score = 0
    if 400 <= metrics['word_count'] <= 800:
        format_score += 40
    elif 200 <= metrics['word_count'] <= 900:
        format_score += 20
        
    format_score += 20 if metrics['email'] else 0
    format_score += 10 if metrics['phone'] else 0
    format_score += 10 if metrics['sections']['education'] else 0
    format_score += 10 if metrics['sections']['experience'] else 0
    format_score += 10 if metrics['sections']['projects'] else 0
    
    format_score = min(100, format_score)
    score += (format_score * 0.2)
    
    metrics['ats_score'] = min(100, round(score, 1))
    
    # Radar Chart Data
    metrics['radar_data'] = [
        semantic_score,           # Semantics
        format_score,             # Format Health
        impact_score,             # Impact/Metrics
        (metrics['action_verbs_found']/10)*100, # Action Verbs
        keyword_match_pct         # Keyword Match
    ]
    
    # 6. Generate Recruiter Actionable Feedback
    tips = []
    if metrics['word_count'] < 300:
        tips.append(("⚠️ Low Word Count", f"Your resume has {metrics['word_count']} words. ATS algorithms highly prefer resumes between 400-800 words to ensure enough keyword density."))
    if metrics['quantifiable_metrics'] == 0:
        tips.append(("⚠️ Zero Quantifiable Metrics", "Recruiters look for data-driven results! We didn't find any percentages (%), dollar signs ($), or numerical scales. Try converting 'Improved sales' to 'Improved sales by 30%'."))
    if len(missing) > 3:
        tips.append(("❌ Critical Keyword Gap", f"You are missing {len(missing)} core skills explicitly requested in the Job Description. The ATS will auto-reject this resume unless customized further."))
    if metrics['action_verbs_found'] < 5:
        tips.append(("📝 Weak Action Verbs", "Avoid passive language like 'Responsible for'. Start your bullet points with power verbs like 'Architected', 'Spearheaded', or 'Maximized'."))
    if not metrics['sections']['experience']:
        tips.append(("🔧 Missing Experience Section", "The ATS could not cleanly detect a standard 'Experience' or 'Employment' header block."))
        
    if not tips:
        tips.append(("✅ Excellent Formatting!", "Your resume fits all major ATS technical constraints perfectly!"))
        
    metrics['feedback_tips'] = tips
    
    return metrics
